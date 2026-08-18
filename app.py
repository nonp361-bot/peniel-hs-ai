import streamlit as st
import google.generativeai as genai
import pypdf
import json
import os
import re
import urllib.request
from io import BytesIO

# --- Word (DOCX) 생성 라이브러리 예외 처리 ---
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# --- ReportLab PDF 라이브러리 연동 ---
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

# --- 1. 페이지 기본 설정 ---
st.set_page_config(
    page_title="브니엘고 AI 생기부 정밀 평가 시스템",
    page_icon="🎓",
    layout="wide"
)

# --- 2. 로컬 채점 기준 보관 폴더 (영구 저장) ---
CRITERIA_DB_DIR = "criteria_database"
os.makedirs(CRITERIA_DB_DIR, exist_ok=True)

# --- 3. 헬퍼 함수 ---
def extract_text_from_pdf_stream(pdf_file):
    """RAM(메모리) 상에서 직접 PDF 텍스트 추출 및 정제 - 개인정보 휘발성 처리"""
    if pdf_file is None:
        return ""
    try:
        pdf_reader = pypdf.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            extracted = page.extract_text()
            if extracted:
                cleaned = extracted.replace("\r", "\n")
                text += cleaned + "\n"
        return text
    except Exception as e:
        st.error(f"PDF 텍스트 추출 중 오류: {e}")
        return ""

def load_local_file_text(filename):
    """저장된 채점 기준 파일 텍스트 로드"""
    path = os.path.join(CRITERIA_DB_DIR, filename)
    if not os.path.exists(path):
        return ""
    if filename.lower().endswith(".pdf"):
        try:
            reader = pypdf.PdfReader(path)
            text = ""
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
            return text
        except Exception:
            return ""
    else:
        try:
            with open(path, "r", encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            try:
                with open(path, "r", encoding="cp949") as f:
                    return f.read()
            except Exception:
                return ""

def get_korean_font_path():
    local_font_name = "NanumGothic.ttf"
    if os.path.exists(local_font_name):
        return local_font_name
    paths = [
        "C:\\Windows\\Fonts\\malgun.ttf",
        "/System/Library/Fonts/Supplemental/AppleGothic.ttf",
        "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
    ]
    for p in paths:
        if os.path.exists(p):
            return p
    font_url = "https://github.com/google/fonts/raw/main/ofl/nanumgothic/NanumGothic-Regular.ttf"
    try:
        urllib.request.urlretrieve(font_url, local_font_name)
        if os.path.exists(local_font_name):
            return local_font_name
    except Exception:
        pass
    return None

def normalize_feedback_text(text):
    """AI 응답이 프롬프트 지시(줄바꿈 포함)를 따르지 않고 줄바꿈 없는 한 덩어리 문단으로
    돌아오는 경우를 대비한 코드 단 안전장치. 이미 줄바꿈이 충분히 들어있어 잘 구조화된
    텍스트는 그대로 두고, 그렇지 않은 경우에만 항목별 소제목/번호목록/수정 전·후 패턴
    앞에 줄바꿈을 강제로 삽입하여 '기재 우수 사항'처럼 문단이 나뉘어 보이도록 보정한다."""
    if not text:
        return text
    t = str(text)

    lines = [ln for ln in t.split("\n") if ln.strip()]
    avg_len = sum(len(ln) for ln in lines) / len(lines) if lines else 0
    # 이미 충분히 여러 줄로 잘 구조화된 경우(평균 한 줄 길이가 짧음)에는 그대로 반환
    if len(lines) >= 3 and avg_len < 150:
        return t

    # 1) [2번. 교사 관찰 결여] 형태의 항목 소제목 앞에 빈 줄 + 줄바꿈 삽입
    t = re.sub(r"\s*(\[[^\[\]]{2,40}\])", r"\n\n\1", t)
    # 2) "1. (과목명)" 또는 "1. (부원명)" 처럼 번호+괄호로 시작하는 문제 문장 항목 앞에 줄바꿈 삽입
    t = re.sub(r"(?<!\n)\s*(\d{1,2}[\.\)]\s*\([^)]{1,25}\))", r"\n\1", t)
    # 3) 괄호 태그 없이 "1. '문장'" 형태로 오는 번호 목록도 줄바꿈 삽입
    t = re.sub(r"(?<!\n)\s*(\d{1,2}[\.\)]\s*['\"'])", r"\n\1", t)
    # 4) "- 수정 전:" / "- 수정 후:" / "수정 전:" / "수정 후:" 앞에 줄바꿈 삽입
    t = re.sub(r"\s*(-?\s*수정\s*(전|후)\s*[:：])", r"\n\1", t)
    # 5) 문장 종결 따옴표(') 뒤에 공백 없이 바로 다음 번호 항목이 이어 붙는 경우 분리
    t = re.sub(r"(['’」\)])\s*(?=\d{1,2}[\.\)])", r"\1\n", t)
    # 6) 연속된 빈 줄은 최대 1개까지만 허용하도록 정리
    t = re.sub(r"\n{3,}", "\n\n", t)

    return t.strip()

# --- 세션 상태 초기화 (초기 미선택 고정) ---
if "feedback_main_cat" not in st.session_state:
    st.session_state["feedback_main_cat"] = "--- 피드백 대분류 선택 ---"

if "feedback_sub_cat" not in st.session_state:
    st.session_state["feedback_sub_cat"] = "--- 세부 영역 선택 ---"

# --- 4. 사이드바 구성 ---
st.sidebar.markdown("<h1 style='text-align: center;'>🎓</h1>", unsafe_allow_html=True)
st.sidebar.markdown("<h3 style='text-align: center;'>브니엘고 AI 평가 시스템</h3>", unsafe_allow_html=True)
st.sidebar.divider()

# 4-1. 입학사정관 선택
st.sidebar.markdown("### 🎯 1. 입학사정관 유형 선택")
evaluator_mode = st.sidebar.radio(
    "사정관 관점을 선택하세요.",
    ["인서울 입학사정관", "지거국 입학사정관"],
    index=0,
    key="evaluator_mode_radio"
)

st.sidebar.divider()

# 4-2. 피드백 버전 선택 (초기 미선택 상태 고정)
st.sidebar.markdown("### 📝 2. 평가 및 피드백 버전 선택")

def on_main_cat_change():
    st.session_state["feedback_sub_cat"] = "--- 세부 영역 선택 ---"
    if "eval_result" in st.session_state:
        del st.session_state["eval_result"]

feedback_category = st.sidebar.selectbox(
    "피드백 대분류를 선택하세요 (필수)",
    [
        "--- 피드백 대분류 선택 ---",
        "교사전용 피드백 버전",
        "학생용 피드백 버전"
    ],
    key="feedback_main_cat",
    on_change=on_main_cat_change
)

selected_feedback_type = "미선택"

if feedback_category == "교사전용 피드백 버전":
    selected_feedback_type = st.sidebar.radio(
        "세부 평가 영역 선택 (필수)",
        [
            "--- 세부 영역 선택 ---",
            "과목세부능력 특기사항 전용 피드백",
            "동아리 특기사항 전용 피드백",
            "생기부 종합 전용 피드백"
        ],
        key="feedback_sub_cat"
    )
elif feedback_category == "학생용 피드백 버전":
    selected_feedback_type = "학생전용 피드백"
    st.sidebar.info("🎓 학생 1인에 대한 냉정한 현위치 진단, 강점/약점 분석 및 앞으로의 탐구 솔루션을 제공합니다.")

st.sidebar.divider()

# 4-3. API 키 설정
api_key = ""
if os.path.exists("api_key.txt"):
    try:
        with open("api_key.txt", "r", encoding="utf-8") as f:
            api_key = f.read().strip()
    except Exception:
        pass

if not api_key and "GEMINI_API_KEY" in st.secrets:
    api_key = st.secrets["GEMINI_API_KEY"]

if not api_key:
    api_key = st.sidebar.text_input("🔑 구글 Gemini API Key 입력", type="password", key="api_key_input")
    if api_key:
        genai.configure(api_key=api_key)
        st.sidebar.success("✅ 인증 완료!")
else:
    genai.configure(api_key=api_key)
    st.sidebar.success("🔑 API 키 자동 인증 완료")

st.sidebar.divider()

# 4-4. 모델 선택
model_option = st.sidebar.selectbox(
    "🤖 Gemini AI 모델 선택",
    ["gemini-2.5-flash", "gemini-1.5-pro", "gemini-1.5-flash"],
    index=0,
    key="model_option_select"
)

st.sidebar.divider()

# 4-5. 채점 기준 파일 관리
st.sidebar.markdown("### 📚 대학별 채점기준 DB (영구보관)")
uploaded_criteria = st.sidebar.file_uploader(
    "채점기준 PDF/TXT 업로드 (무제한)",
    type=["pdf", "txt"],
    accept_multiple_files=True,
    key="criteria_uploader"
)

if uploaded_criteria:
    for file in uploaded_criteria:
        dest_path = os.path.join(CRITERIA_DB_DIR, file.name)
        with open(dest_path, "wb") as f:
            f.write(file.getbuffer())
    st.sidebar.success(f"💾 {len(uploaded_criteria)}개 파일 DB 누적 저장 완료!")
    st.rerun()

accumulated_files = os.listdir(CRITERIA_DB_DIR)
selected_criteria_files = []

if accumulated_files:
    st.sidebar.markdown("**📌 반영할 채점 기준 선택**")
    for file_name in accumulated_files:
        if st.sidebar.checkbox(file_name, value=True, key=f"check_{file_name}"):
            selected_criteria_files.append(file_name)
    
    if st.sidebar.button("🗑️ 선택 파일 DB에서 삭제", type="secondary", key="delete_criteria_btn"):
        for file_name in accumulated_files:
            if st.session_state.get(f"check_{file_name}", False):
                os.remove(os.path.join(CRITERIA_DB_DIR, file_name))
        st.rerun()

# --- 5. 메인 화면 헤더 및 실시간 채점 기준표 노출 ---
st.title("🏫 브니엘고등학교 AI 생기부 정밀 평가 시스템")

is_unselected = ("선택" in selected_feedback_type) or (selected_feedback_type == "미선택")
display_mode_str = selected_feedback_type if not is_unselected else "영역 미선택 (좌측 사이드바에서 선택)"

st.markdown(f"**현재 관점:** `{evaluator_mode}` | **선택된 피드백 모드:** `<font color='#1E3A8A'><b>{display_mode_str}</b></font>` | **적용 기준 파일:** `{len(selected_criteria_files)}개`", unsafe_allow_html=True)

st.markdown("### 📋 AI 실시간 통합 채점 기준표 (메인 상시 노출)")

# 📌 [경로 A] 피드백 영역 미선택 시
if is_unselected:
    st.warning("👈 **왼쪽 사이드바의 [2. 평가 및 피드백 버전 선택]에서 피드백 대분류와 세부 평가 영역을 먼저 선택해 주세요.**")

# 📌 [경로 B-1] 과세특 7대 점검 항목 표 노출
elif selected_feedback_type == "과목세부능력 특기사항 전용 피드백":
    st.info("💡 **과목 세부능력 특기사항 교사 자가점검 7대 핵심 채점기준표 (100점 만점)**")
    st.markdown("""
    <div style="background-color: #F8FAFC; border: 1.5px solid #CBD5E1; border-radius: 8px; padding: 15px; margin-bottom: 15px;">
        <table style="width: 100%; border-collapse: collapse; text-align: left; font-size: 14.5px;">
            <thead>
                <tr style="background-color: #1E3A8A; color: white;">
                    <th style="padding: 10px; border: 1px solid #CBD5E1; text-align: center; width: 60px;">번호</th>
                    <th style="padding: 10px; border: 1px solid #CBD5E1; width: 340px;">과세특 교사 7대 점검 항목</th>
                    <th style="padding: 10px; border: 1px solid #CBD5E1; text-align: center; width: 80px;">배점</th>
                    <th style="padding: 10px; border: 1px solid #CBD5E1;">핵심 점검 내용 및 세부 가이드</th>
                </tr>
            </thead>
            <tbody>
                <tr style="background-color: #EFF6FF;">
                    <td style="padding: 10px; border: 1px solid #CBD5E1; font-weight: bold; text-align: center;">1</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1; font-weight: bold;">학생의 교과적 역량을 잘 보여주는 기록인가</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1; text-align: center; font-weight: bold; color: #DC2626;">20점</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1;">단순 수업 참여를 넘어 교과 개념의 깊이 있는 이해와 지적 도약이 드러나는가 (수학, 영어, 통합사회/과학 등 전 과목 포함)</td>
                </tr>
                <tr style="background-color: #EFF6FF;">
                    <td style="padding: 10px; border: 1px solid #CBD5E1; font-weight: bold; text-align: center;">2</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1; font-weight: bold;">교사의 관찰이 들어간 기록인가</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1; text-align: center; font-weight: bold; color: #DC2626;">20점</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1;">학생 보고서 요약이 아닌 수업 중 교사가 직접 관찰한 행동과 질문이 기재되었는가</td>
                </tr>
                <tr style="background-color: #EFF6FF;">
                    <td style="padding: 10px; border: 1px solid #CBD5E1; font-weight: bold; text-align: center;">3</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1; font-weight: bold;">교과의 역량이 들어갔는가</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1; text-align: center; font-weight: bold; color: #DC2626;">20점</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1;">해당 교과목 고유의 핵심 성취기준 및 사고방식이 반영되었는가</td>
                </tr>
                <tr style="background-color: #FEF3C7;">
                    <td style="padding: 10px; border: 1px solid #CBD5E1; font-weight: bold; text-align: center;">4</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1; font-weight: bold;">학생 간 복붙한 기록이 없는가</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1; text-align: center; font-weight: bold; color: #DC2626;">10점</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1;">수강생 전체/일부 간 동일한 서술어나 표준화된 문장이 반복적으로 쓰이지 않았는가</td>
                </tr>
                <tr style="background-color: #FEF3C7;">
                    <td style="padding: 10px; border: 1px solid #CBD5E1; font-weight: bold; text-align: center;">5</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1; font-weight: bold;">AI를 너무 돌려 맥락에 맞지 않는 단어나 문장이 들어가지 않았는가</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1; text-align: center; font-weight: bold; color: #DC2626;">10점</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1;">구체적 사례 없이 거대 담론이나 어색한 AI 생성 표현만 나열되어 있지 않은가</td>
                </tr>
                <tr style="background-color: #F3F4F6;">
                    <td style="padding: 10px; border: 1px solid #CBD5E1; font-weight: bold; text-align: center;">6</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1; font-weight: bold;">가독성이 높은가</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1; text-align: center; font-weight: bold; color: #DC2626;">10점</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1;">동기-과정-성장의 논리적 구조를 갖추고 문장의 호응과 전달력이 완벽한가</td>
                </tr>
                <tr style="background-color: #F3F4F6;">
                    <td style="padding: 10px; border: 1px solid #CBD5E1; font-weight: bold; text-align: center;">7</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1; font-weight: bold;">생기부 기재 금지 사항이 잘 반영되었는가</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1; text-align: center; font-weight: bold; color: #DC2626;">10점</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1;">대학명, 기관명, 상호명, 강사명 등 기재 불가능한 항목이 철저히 배제되었는가</td>
                </tr>
            </tbody>
        </table>
    </div>
    """, unsafe_allow_html=True)

# 📌 [경로 B-2] 동아리 8대 점검 항목 표 노출
elif selected_feedback_type == "동아리 특기사항 전용 피드백":
    st.info("💡 **동아리 특기사항 교사 자가점검 8대 핵심 채점기준표 (100점 만점)**")
    st.markdown("""
    <div style="background-color: #F8FAFC; border: 1.5px solid #CBD5E1; border-radius: 8px; padding: 15px; margin-bottom: 15px;">
        <table style="width: 100%; border-collapse: collapse; text-align: left; font-size: 14.5px;">
            <thead>
                <tr style="background-color: #1E3A8A; color: white;">
                    <th style="padding: 10px; border: 1px solid #CBD5E1; text-align: center; width: 60px;">번호</th>
                    <th style="padding: 10px; border: 1px solid #CBD5E1; width: 340px;">동아리 교사 8대 점검 항목</th>
                    <th style="padding: 10px; border: 1px solid #CBD5E1; text-align: center; width: 80px;">배점</th>
                    <th style="padding: 10px; border: 1px solid #CBD5E1;">핵심 점검 내용 및 세부 가이드</th>
                </tr>
            </thead>
            <tbody>
                <tr style="background-color: #EFF6FF;">
                    <td style="padding: 10px; border: 1px solid #CBD5E1; font-weight: bold; text-align: center;">1</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1; font-weight: bold;">학생의 학문적 탐구 역량을 잘 보여주는 기록인가</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1; text-align: center; font-weight: bold; color: #DC2626;">15점</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1;">동아리 활동 속에서 심화된 지적 호기심과 학문적 깊이가 드러나는가</td>
                </tr>
                <tr style="background-color: #EFF6FF;">
                    <td style="padding: 10px; border: 1px solid #CBD5E1; font-weight: bold; text-align: center;">2</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1; font-weight: bold;">교사의 관찰이 들어간 기록인가</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1; text-align: center; font-weight: bold; color: #DC2626;">15점</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1;">단순 동아리 소개나 활동 나열이 아닌 지도교사의 구체적 관찰 사실이 기재되었는가</td>
                </tr>
                <tr style="background-color: #EFF6FF;">
                    <td style="padding: 10px; border: 1px solid #CBD5E1; font-weight: bold; text-align: center;">3</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1; font-weight: bold;">학생이 이수한 교과목에서 심화 탐구한 내용이 들어갔는가</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1; text-align: center; font-weight: bold; color: #DC2626;">15점</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1;">정규 교과 지식을 동아리 활동 속에서 어떻게 확장하고 심화했는지 연결고리가 있는가</td>
                </tr>
                <tr style="background-color: #FEF3C7;">
                    <td style="padding: 10px; border: 1px solid #CBD5E1; font-weight: bold; text-align: center;">4</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1; font-weight: bold;">학생 간 복붙한 기록이 없는가</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1; text-align: center; font-weight: bold; color: #DC2626;">10점</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1;">동아리 부원 전원에게 동일하게 적용된 복사-붙여넣기식 서술이 배제되었는가</td>
                </tr>
                <tr style="background-color: #FEF3C7;">
                    <td style="padding: 10px; border: 1px solid #CBD5E1; font-weight: bold; text-align: center;">5</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1; font-weight: bold;">AI를 너무 돌려 맥락에 맞지 않는 단어나 문장이 들어가지 않았는가</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1; text-align: center; font-weight: bold; color: #DC2626;">10점</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1;">개인별 구체적 역할 없이 거대 담론이나 어색한 AI 생성 표현만 가득하지 않은가</td>
                </tr>
                <tr style="background-color: #F3F4F6;">
                    <td style="padding: 10px; border: 1px solid #CBD5E1; font-weight: bold; text-align: center;">6</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1; font-weight: bold;">가독성이 높은가</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1; text-align: center; font-weight: bold; color: #DC2626;">15점</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1;">주도적 동기, 협업 과정, 성장의 서사가 논리적으로 매끄럽게 연결되는가</td>
                </tr>
                <tr style="background-color: #F3F4F6;">
                    <td style="padding: 10px; border: 1px solid #CBD5E1; font-weight: bold; text-align: center;">7</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1; font-weight: bold;">생기부 기재 금지 사항이 잘 반영되어 있는가</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1; text-align: center; font-weight: bold; color: #DC2626;">10점</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1;">외부 기관명, 대회 명칭, 인증 시험명 등 기재 불가 항목이 철저히 배제되었는가</td>
                </tr>
                <tr style="background-color: #F3F4F6;">
                    <td style="padding: 10px; border: 1px solid #CBD5E1; font-weight: bold; text-align: center;">8</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1; font-weight: bold;">오탈자가 없는가</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1; text-align: center; font-weight: bold; color: #DC2626;">10점</td>
                    <td style="padding: 10px; border: 1px solid #CBD5E1;">맞춤법, 띄어쓰기, 문장 부호 등의 오탈자가 완벽하게 검수되었는가</td>
                </tr>
            </tbody>
        </table>
    </div>
    """, unsafe_allow_html=True)

# 📌 [경로 C] 기타 생기부 영역 및 학생용 탭 선택 시 삼분 배점 기준표 표출
else:
    col_c1, col_c2, col_c3 = st.columns(3)
    with col_c1:
        st.info("📕 **1. 학업역량 (40점 만점)**")
        st.markdown("- **성취도 분포 및 이수환경의 상대적 우위성**: 공통과목(공통수학, 영어 등) 및 선택과목 종합 해석")
        st.markdown("- **행동 동기 및 어려움 극복 서사 기반 학업태도**: 교사 직접 관찰 기반 자발적 탐구 열의")
        st.markdown("- **디지털 리터러시 및 비판적 미디어 탐구 역량**: Bloom 5-6단계 사고 및 비판적 미디어 활용")
    with col_c2:
        st.success("📗 **2. 진로역량 (40점 만점)**")
        st.markdown("- **전공 연계 교과의 위계적 이수 노력**: 권장 이수과목, 과목 위계성 및 선택 동기")
        st.markdown("- **전공 관련 주요 교과 성취도 차별성**: 전공 관련 교과 성취도 차별성 및 전공적 사고")
        st.markdown("- **교과-창체 연계 진로 에피소드**: 문헌 비판적 독해 및 활동 간 수직/수평적 일관성")
    with col_c3:
        st.warning("📘 **3. 공동체역량 (20점 만점)**")
        st.markdown("- **다원적 환경에서의 협업 및 소통 역량**: 실질적 역할 기여 및 갈등/오해 조율")
        st.markdown("- **특정 대상을 도운 구체적 나눔과 배려**: 구체적 에피소드 중심의 사회적 가치 실천")
        st.markdown("- **성실성과 규칙 준수 및 자발적 리더십**: 무단 출결 배제, 규칙 준수 및 솔선수범")

st.divider()

# --- 6. PDF 및 DOCX 보고서 생성 함수 ---
def generate_pdf_report(eval_data, student_filename, mode, fb_type):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=35, leftMargin=35, topMargin=35, bottomMargin=35)
    
    font_path = get_korean_font_path()
    font_name = "Helvetica"
    if font_path:
        try:
            pdfmetrics.registerFont(TTFont("KoreanFont", font_path))
            font_name = "KoreanFont"
        except Exception:
            pass

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('DocTitle', parent=styles['Heading1'], fontName=font_name, fontSize=15, leading=19, textColor=colors.HexColor('#1E3A8A'), alignment=1, spaceAfter=8)
    subtitle_style = ParagraphStyle('DocSubtitle', parent=styles['Normal'], fontName=font_name, fontSize=8.5, leading=11, textColor=colors.HexColor('#4B5563'), alignment=1, spaceAfter=12)
    h1_style = ParagraphStyle('H1', parent=styles['Heading2'], fontName=font_name, fontSize=11, leading=14, textColor=colors.HexColor('#1E3A8A'), spaceBefore=8, spaceAfter=4, keepWithNext=True)
    body_style = ParagraphStyle('Body', parent=styles['Normal'], fontName=font_name, fontSize=8, leading=11, textColor=colors.HexColor('#1F2937'), spaceAfter=3)
    
    def create_box(text, bg_hex, border_hex, text_hex):
        inner_style = ParagraphStyle('BoxInner', fontName=font_name, fontSize=7.5, leading=10, textColor=colors.HexColor(text_hex))
        p = Paragraph(text, inner_style)
        t = Table([[p]], colWidths=[520])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,-1), colors.HexColor(bg_hex)),
            ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor(border_hex)),
            ('TOPPADDING', (0,0), (-1,-1), 5), ('BOTTOMPADDING', (0,0), (-1,-1), 5),
            ('LEFTPADDING', (0,0), (-1,-1), 6), ('RIGHTPADDING', (0,0), (-1,-1), 6),
        ]))
        return t

    def markdown_to_flowables(text, text_hex):
        """AI가 생성한 마크다운풍 텍스트(번호목록/불릿/굵게)를 리포트랩 문단 리스트로 변환하여
        PDF 안에서도 화면(st.markdown)과 동일하게 문단·목록이 구분되어 보이도록 처리

        [중요 - LayoutError 근본 방지책]
        AI 응답에 줄바꿈(\\n)이 제대로 들어있지 않고 거대한 한 덩어리 문단으로 오는 경우가 있다.
        이 경우 문단 1개가 그대로 표의 행(row) 1개가 되어, 페이지 한 장보다 길면 다시 LayoutError가 난다.
        따라서 텍스트의 줄바꿈 여부와 무관하게, 일정 글자 수(MAX_CHARS)를 넘는 덩어리는
        공백 단위로 강제로 잘게 쪼개어 여러 개의 짧은 문단(=여러 행)으로 만든다.
        이렇게 하면 어떤 경우에도 개별 행이 한 페이지보다 커질 수 없다.
        """
        MAX_CHARS = 220

        if not text:
            return [Paragraph("-", ParagraphStyle(
                'mdEmpty', fontName=font_name, fontSize=7.5, leading=10.5,
                textColor=colors.HexColor(text_hex)
            ))]

        normal_style = ParagraphStyle(
            'mdNormal', fontName=font_name, fontSize=7.5, leading=11,
            textColor=colors.HexColor(text_hex), spaceAfter=5
        )
        continuation_style = ParagraphStyle(
            'mdContinuation', fontName=font_name, fontSize=7.5, leading=11,
            textColor=colors.HexColor(text_hex), leftIndent=10, spaceAfter=5
        )
        section_style = ParagraphStyle(
            'mdSection', fontName=font_name, fontSize=8.3, leading=11.5,
            textColor=colors.HexColor(text_hex), spaceBefore=9, spaceAfter=4
        )
        numbered_style = ParagraphStyle(
            'mdNumbered', fontName=font_name, fontSize=7.8, leading=11.2,
            textColor=colors.HexColor(text_hex), spaceBefore=7, spaceAfter=2
        )
        numbered_continuation_style = ParagraphStyle(
            'mdNumberedContinuation', fontName=font_name, fontSize=7.8, leading=11.2,
            textColor=colors.HexColor(text_hex), leftIndent=12, spaceAfter=2
        )
        bullet_style = ParagraphStyle(
            'mdBullet', fontName=font_name, fontSize=7.5, leading=10.8,
            textColor=colors.HexColor(text_hex), leftIndent=16, spaceAfter=4
        )
        bullet_continuation_style = ParagraphStyle(
            'mdBulletContinuation', fontName=font_name, fontSize=7.5, leading=10.8,
            textColor=colors.HexColor(text_hex), leftIndent=26, spaceAfter=4
        )
        sub_bullet_style = ParagraphStyle(
            'mdSubBullet', fontName=font_name, fontSize=7.3, leading=10.5,
            textColor=colors.HexColor(text_hex), leftIndent=30, spaceAfter=4
        )

        def chunk_long_text(raw, max_chars=MAX_CHARS):
            """공백 단위로 텍스트를 max_chars 이하 조각으로 강제 분할.
            공백이 전혀 없는 비정상적으로 긴 토큰(예: 끊어지지 않는 긴 문자열)도
            안전하게 강제 슬라이싱하여 절대 max_chars*2를 넘지 않도록 보장한다."""
            raw = raw.strip()
            if not raw:
                return []
            if len(raw) <= max_chars:
                return [raw]

            words = raw.split(" ")
            chunks, current = [], ""
            for w in words:
                # 공백이 없는 단일 토큰 자체가 너무 길면 강제로 잘라냄
                while len(w) > max_chars:
                    if current:
                        chunks.append(current)
                        current = ""
                    chunks.append(w[:max_chars])
                    w = w[max_chars:]
                candidate = f"{current} {w}".strip() if current else w
                if len(candidate) > max_chars and current:
                    chunks.append(current)
                    current = w
                else:
                    current = candidate
            if current:
                chunks.append(current)
            return chunks

        def emit(raw_content, prefix, first_style, continuation_style_):
            """prefix(번호/불릿 등)는 첫 조각에만 붙이고, 나머지 조각은 이어지는 문단으로 처리"""
            chunks = chunk_long_text(raw_content)
            if not chunks:
                return
            for idx, chunk in enumerate(chunks):
                chunk_fmt = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", chunk)
                if idx == 0:
                    flowables.append(Paragraph(f"{prefix}{chunk_fmt}", first_style))
                else:
                    flowables.append(Paragraph(chunk_fmt, continuation_style_))

        flowables = []
        for raw_line in str(text).split("\n"):
            line = raw_line.strip()
            if not line:
                continue
            leading_spaces = len(raw_line) - len(raw_line.lstrip(" "))

            # [2번. 교사 관찰 결여] 형태의 항목별 소제목 인식 (굵게 변환 이전의 원문 기준으로 판별)
            m_section = re.match(r"^\**\[(.+?)\]\**\s*(.*)", line)
            if m_section:
                title_fmt = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", m_section.group(1))
                flowables.append(Paragraph(f"▍<b>{title_fmt}</b>", section_style))
                if m_section.group(2):
                    emit(m_section.group(2), "", normal_style, continuation_style)
                continue

            m_num = re.match(r"^(\d+)[\.\)]\s+(.*)", line)
            m_bullet = re.match(r"^[-•*○ㅇ]\s+(.*)", line)

            if m_num:
                emit(m_num.group(2), f"<b>{m_num.group(1)}.</b> ", numbered_style, numbered_continuation_style)
            elif m_bullet:
                style = sub_bullet_style if leading_spaces >= 2 else bullet_style
                emit(m_bullet.group(1), "• ", style, bullet_continuation_style)
            else:
                emit(line, "", normal_style, continuation_style)
        return flowables


    def create_feedback_box(title, text, bg_hex, border_hex, text_hex):
        """제목 + 좌측 컬러 포인트 바 + 문단/목록이 구분된 피드백 카드를 생성
        (사이트의 st.info/success/warning/error 색상 구분을 PDF에서도 동일하게 재현)

        [중요] 문단 리스트를 표 셀 하나에 몰아넣지 않고, '문단 1개 = 표의 행(row) 1개'
        구조로 만들어야 ReportLab이 페이지 경계에서 행 단위로 자연스럽게 쪼갤 수 있다.
        (전수 점검 기능으로 피드백 분량이 길어지면서, 셀 하나에 내용을 몰아넣던 기존 방식은
        한 페이지 높이를 넘는 순간 LayoutError가 발생했음 — 이를 근본적으로 해결하기 위한 구조)
        """
        title_style = ParagraphStyle(
            'BoxTitle', fontName=font_name, fontSize=9, leading=12,
            textColor=colors.HexColor(text_hex), spaceAfter=0
        )
        body_flowables = markdown_to_flowables(text, text_hex)

        rows = [[Paragraph(f"<b>{title}</b>", title_style)]]
        for fl in body_flowables:
            rows.append([fl])

        t = Table(rows, colWidths=[520])
        style_cmds = [
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor(bg_hex)),
            ('BOX', (0, 0), (-1, -1), 0.75, colors.HexColor(border_hex)),
            ('LINEBEFORE', (0, 0), (0, -1), 3, colors.HexColor(border_hex)),
            ('TOPPADDING', (0, 0), (-1, -1), 1.5), ('BOTTOMPADDING', (0, 0), (-1, -1), 1.5),
            ('LEFTPADDING', (0, 0), (-1, -1), 10), ('RIGHTPADDING', (0, 0), (-1, -1), 9),
            # 제목 행 상단 여백, 마지막 행 하단 여백만 넉넉하게
            ('TOPPADDING', (0, 0), (0, 0), 7),
            ('BOTTOMPADDING', (0, -1), (0, -1), 7),
        ]
        t.setStyle(TableStyle(style_cmds))
        # 표가 페이지 경계에서 행 단위로 자연스럽게 분할되도록 명시적으로 허용
        t.splitByRow = 1
        t.hAlign = 'LEFT'
        return t

    story = []
    story.append(Paragraph(f"🎓 브니엘고 AI 생기부 평가 리포트 ({mode})", title_style))
    story.append(Paragraph(f"대상 파일: {student_filename}  |  평가 모드: {fb_type}", subtitle_style))
    story.append(Spacer(1, 4))
    
    if fb_type == "과목세부능력 특기사항 전용 피드백" and "setuk_eval" in eval_data:
        st_data = eval_data["setuk_eval"]
        t_scores = st_data.get("scores", {})
        table_data = [
            [Paragraph("<b>과세특 교사 7대 점검 항목</b>", body_style), Paragraph("<b>배점</b>", body_style), Paragraph("<b>획득 점수</b>", body_style)],
            [Paragraph("1. 학생의 교과적 역량 서술", body_style), Paragraph("20점", body_style), Paragraph(f"{t_scores.get('academic_competence', 0)}점", body_style)],
            [Paragraph("2. 교사의 직접 관찰 반영", body_style), Paragraph("20점", body_style), Paragraph(f"{t_scores.get('teacher_observation', 0)}점", body_style)],
            [Paragraph("3. 교과의 핵심 역량 기재", body_style), Paragraph("20점", body_style), Paragraph(f"{t_scores.get('subject_competence', 0)}점", body_style)],
            [Paragraph("4. 학생 간 복붙 기재 여부", body_style), Paragraph("10점", body_style), Paragraph(f"{t_scores.get('duplication', 0)}점", body_style)],
            [Paragraph("5. AI 대필 / 문맥 어색함 검증", body_style), Paragraph("10점", body_style), Paragraph(f"{t_scores.get('ai_overuse', 0)}점", body_style)],
            [Paragraph("6. 가독성 및 문장 구조", body_style), Paragraph("10점", body_style), Paragraph(f"{t_scores.get('readability', 0)}점", body_style)],
            [Paragraph("7. 생기부 기재 금지 사항 준수", body_style), Paragraph("10점", body_style), Paragraph(f"{t_scores.get('prohibited_items', 0)}점", body_style)],
            [Paragraph("<b>✨ 최종 과세특 작성 품질 총점</b>", body_style), Paragraph("<b>100점</b>", body_style), Paragraph(f"<b><font color='#EF4444'>{t_scores.get('total', 0)} / 100</font></b>", body_style)]
        ]
        t_score = Table(table_data, colWidths=[220, 100, 200])
        t_score.setStyle(TableStyle([('BACKGROUND', (0,0), (-1,0), colors.HexColor('#F3F4F6')), ('ALIGN', (0,0), (-1,-1), 'CENTER'), ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#E5E7EB')), ('BACKGROUND', (0,8), (-1,8), colors.HexColor('#FEF2F2'))]))
        story.append(t_score)
        story.append(Spacer(1, 10))
        story.append(Paragraph("👩‍🏫 [과세특 개선 피드백 리포트]", h1_style))
        story.append(Spacer(1, 3))
        story.append(create_feedback_box("📝 종합 총평", st_data.get('overall_summary', ''), '#EFF6FF', '#3B82F6', '#1E3A8A'))
        story.append(Spacer(1, 7))
        story.append(create_feedback_box("👍 기재 우수 사항", st_data.get('good_points', ''), '#F0FDF4', '#22C55E', '#14532D'))
        story.append(Spacer(1, 7))
        story.append(create_feedback_box("⚠️ 보완 및 수정 필요사항", st_data.get('improvements', ''), '#FFFBEB', '#F59E0B', '#78350F'))
        story.append(Spacer(1, 7))
        story.append(create_feedback_box("✏️ 수정·보완 추천 문장 예시", st_data.get('revision_examples', ''), '#FEF2F2', '#EF4444', '#991B1B'))
    elif fb_type == "동아리 특기사항 전용 피드백" and "club_eval" in eval_data:
        c_data = eval_data["club_eval"]
        c_scores = c_data.get("scores", {})
        table_data = [
            [Paragraph("<b>동아리 교사 8대 점검 항목</b>", body_style), Paragraph("<b>배점</b>", body_style), Paragraph("<b>획득 점수</b>", body_style)],
            [Paragraph("1. 학생의 학문적 탐구 역량", body_style), Paragraph("15점", body_style), Paragraph(f"{c_scores.get('academic_inquiry', 0)}점", body_style)],
            [Paragraph("2. 교사의 직접 관찰 반영", body_style), Paragraph("15점", body_style), Paragraph(f"{c_scores.get('teacher_observation', 0)}점", body_style)],
            [Paragraph("3. 교과목 심화 탐구 연계", body_style), Paragraph("15점", body_style), Paragraph(f"{c_scores.get('subject_connection', 0)}점", body_style)],
            [Paragraph("4. 학생 간 복붙 기재 여부", body_style), Paragraph("10점", body_style), Paragraph(f"{c_scores.get('duplication', 0)}점", body_style)],
            [Paragraph("5. AI 대필 / 문맥 어색함 검증", body_style), Paragraph("10점", body_style), Paragraph(f"{c_scores.get('ai_overuse', 0)}점", body_style)],
            [Paragraph("6. 가독성 및 문장 구조", body_style), Paragraph("15점", body_style), Paragraph(f"{c_scores.get('readability', 0)}점", body_style)],
            [Paragraph("7. 생기부 기재 금지 사항 준수", body_style), Paragraph("10점", body_style), Paragraph(f"{c_scores.get('prohibited_items', 0)}점", body_style)],
            [Paragraph("8. 오탈자 검수", body_style), Paragraph("10점", body_style), Paragraph(f"{c_scores.get('typo_check', 0)}점", body_style)],
            [Paragraph("<b>✨ 최종 동아리 작성 품질 총점</b>", body_style), Paragraph("<b>100점</b>", body_style), Paragraph(f"<b><font color='#EF4444'>{c_scores.get('total', 0)} / 100</font></b>", body_style)]
        ]
        t_score = Table(table_data, colWidths=[220, 100, 200])
        t_score.setStyle(TableStyle([('BACKGROUND', (0,0), (-1,0), colors.HexColor('#F3F4F6')), ('ALIGN', (0,0), (-1,-1), 'CENTER'), ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#E5E7EB')), ('BACKGROUND', (0,9), (-1,9), colors.HexColor('#FEF2F2'))]))
        story.append(t_score)
        story.append(Spacer(1, 10))
        story.append(Paragraph("👩‍🏫 [동아리 특기사항 개선 피드백 리포트]", h1_style))
        story.append(Spacer(1, 3))
        story.append(create_feedback_box("📝 종합 총평", c_data.get('overall_summary', ''), '#EFF6FF', '#3B82F6', '#1E3A8A'))
        story.append(Spacer(1, 7))
        story.append(create_feedback_box("👍 기재 우수 사항", c_data.get('good_points', ''), '#F0FDF4', '#22C55E', '#14532D'))
        story.append(Spacer(1, 7))
        story.append(create_feedback_box("⚠️ 보완 및 수정 필요사항", c_data.get('improvements', ''), '#FFFBEB', '#F59E0B', '#78350F'))
        story.append(Spacer(1, 7))
        story.append(create_feedback_box("✏️ 수정·보완 추천 문장 예시", c_data.get('revision_examples', ''), '#FEF2F2', '#EF4444', '#991B1B'))
    else:
        scores = eval_data.get("scores", {})
        table_data = [
            [Paragraph("<b>평가 영역</b>", body_style), Paragraph("<b>반영 요소</b>", body_style), Paragraph("<b>취득 점수</b>", body_style)],
            [Paragraph("I. 학업역량 (40점)", body_style), Paragraph("성취도 분포 / 학업태도 / 비판적 탐구", body_style), Paragraph(f"<b>{scores.get('academic', 0)} / 40</b>", body_style)],
            [Paragraph("II. 진로역량 (40점)", body_style), Paragraph("전공 이수 노력 / 전공 성취도 / 진로 탐색", body_style), Paragraph(f"<b>{scores.get('career', 0)} / 40</b>", body_style)],
            [Paragraph("III. 공동체역량 (20점)", body_style), Paragraph("협업·소통 / 나눔·배려 / 성실성 / 리더십", body_style), Paragraph(f"<b>{scores.get('community', 0)} / 20</b>", body_style)],
            [Paragraph("<b>✨ 최종 환산 총점</b>", body_style), Paragraph("<b>100점 만점 기준 종합 환산점수</b>", body_style), Paragraph(f"<b><font color='#EF4444'>{scores.get('total', 0)} / 100</font></b>", body_style)]
        ]
        t_score = Table(table_data, colWidths=[120, 280, 120])
        t_score.setStyle(TableStyle([('BACKGROUND', (0,0), (-1,0), colors.HexColor('#F3F4F6')), ('ALIGN', (0,0), (-1,-1), 'CENTER'), ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#E5E7EB')), ('BACKGROUND', (0,4), (-1,4), colors.HexColor('#FEF2F2'))]))
        story.append(t_score)
        story.append(Spacer(1, 10))
        if "teacher_feedback" in eval_data:
            story.append(Paragraph(f"👩‍🏫 [교사전용 정밀 피드백: {fb_type}]", h1_style))
            story.append(Spacer(1, 3))
            tf = eval_data["teacher_feedback"]
            story.append(create_feedback_box("👍 장점 분석", tf.get('strength', ''), '#F0FDF4', '#22C55E', '#14532D'))
            story.append(Spacer(1, 7))
            story.append(create_feedback_box("⚠️ 보완점 및 감점 원인", tf.get('weakness', ''), '#FFFBEB', '#F59E0B', '#78350F'))
            if tf.get('quote'):
                story.append(Spacer(1, 7))
                story.append(create_feedback_box("🎯 원문 인용 근거", f"\"{tf['quote']}\"", '#FEF3C7', '#F59E0B', '#451A03'))
        if "student_feedback" in eval_data:
            story.append(Paragraph("🎓 [학생전용 현위치 진단 & 솔루션]", h1_style))
            story.append(Spacer(1, 3))
            sf = eval_data["student_feedback"]
            story.append(create_feedback_box("🔍 입학사정관 관점 냉정한 현위치 진단 (지원 가능 대학 라인)", sf.get("current_position", ""), '#F3F4F6', '#9CA3AF', '#111827'))
            story.append(Spacer(1, 7))
            story.append(create_feedback_box("👍 기존 활동의 주요 강점", sf.get('strength_analysis', ''), '#F0FDF4', '#22C55E', '#14532D'))
            story.append(Spacer(1, 7))
            story.append(create_feedback_box("🚨 치명적인 약점 및 감점 요소", sf.get('weakness_analysis', ''), '#FFFBEB', '#F59E0B', '#78350F'))
            story.append(Spacer(1, 7))
            story.append(create_feedback_box("🚀 향후 보완 추천 활동 및 구체적 탐구 주제 솔루션", sf.get("recommendation", ""), '#FEF2F2', '#EF4444', '#991B1B'))
        
    doc.build(story)
    pdf_bytes = buffer.getvalue()
    buffer.close()
    return pdf_bytes

def generate_docx_report(eval_data, student_filename, mode, fb_type):
    if not DOCX_AVAILABLE:
        return None
    doc = Document()
    doc.add_heading(f"브니엘고 AI 생기부 평가 리포트 ({mode})", 0)
    doc.add_paragraph(f"대상 파일: {student_filename} | 평가 모드: {fb_type}")
    
    if fb_type == "과목세부능력 특기사항 전용 피드백" and "setuk_eval" in eval_data:
        st_data = eval_data["setuk_eval"]
        doc.add_heading("과세특 교사 7대 점검 항목 평가", level=1)
        scores = st_data.get("scores", {})
        doc.add_paragraph(f"최종 과세특 작성 품질 총점: {scores.get('total', 0)} / 100")
        doc.add_heading("종합 평어", level=2)
        doc.add_paragraph(st_data.get("overall_summary", ""))
        doc.add_heading("우수 사항", level=2)
        doc.add_paragraph(st_data.get("good_points", ""))
        doc.add_heading("보완 및 수정 필요사항", level=2)
        doc.add_paragraph(st_data.get("improvements", ""))
        doc.add_heading("수정 추천 문장 예시", level=2)
        doc.add_paragraph(st_data.get("revision_examples", ""))
    elif fb_type == "동아리 특기사항 전용 피드백" and "club_eval" in eval_data:
        c_data = eval_data["club_eval"]
        doc.add_heading("동아리 교사 8대 점검 항목 평가", level=1)
        scores = c_data.get("scores", {})
        doc.add_paragraph(f"최종 동아리 작성 품질 총점: {scores.get('total', 0)} / 100")
        doc.add_heading("종합 평어", level=2)
        doc.add_paragraph(c_data.get("overall_summary", ""))
        doc.add_heading("우수 사항", level=2)
        doc.add_paragraph(c_data.get("good_points", ""))
        doc.add_heading("보완 및 수정 필요사항", level=2)
        doc.add_paragraph(c_data.get("improvements", ""))
        doc.add_heading("수정 추천 문장 예시", level=2)
        doc.add_paragraph(c_data.get("revision_examples", ""))
    else:
        scores = eval_data.get("scores", {})
        doc.add_heading("100점 만점 영역별 스코어 카드", level=1)
        doc.add_paragraph(f"학업역량: {scores.get('academic', 0)} / 40")
        doc.add_paragraph(f"진로역량: {scores.get('career', 0)} / 40")
        doc.add_paragraph(f"공동체역량: {scores.get('community', 0)} / 20")
        doc.add_paragraph(f"최종 환산 총점: {scores.get('total', 0)} / 100")
        
        if "teacher_feedback" in eval_data:
            doc.add_heading(f"교사전용 피드백 ({fb_type})", level=1)
            tf = eval_data["teacher_feedback"]
            doc.add_paragraph(f"장점 분석: {tf.get('strength', '')}")
            doc.add_paragraph(f"보완점 및 감점 원인: {tf.get('weakness', '')}")
            if tf.get('quote'):
                doc.add_paragraph(f"원문 인용 근거: \"{tf['quote']}\"")
                
        if "student_feedback" in eval_data:
            doc.add_heading("학생전용 현위치 진단 & 솔루션", level=1)
            sf = eval_data["student_feedback"]
            doc.add_paragraph(f"현재 위치 진단: {sf.get('current_position', '')}")
            doc.add_paragraph(f"활동 강점: {sf.get('strength_analysis', '')}")
            doc.add_paragraph(f"치명적 약점: {sf.get('weakness_analysis', '')}")
            doc.add_paragraph(f"추천 탐구 주제 솔루션: {sf.get('recommendation', '')}")
            
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    return doc_buffer.getvalue()

# --- 7. 학생부 PDF 제출 및 AI 평가 구역 ---
st.markdown("### 📂 학생부 PDF 제출 및 맞춤형 AI 채점")
st.caption("🔒 제출한 생기부 PDF는 디스크에 저장되지 않으며, RAM에서 읽어 평가 후 즉시 영구 휘발 삭제됩니다.")

student_file = st.file_uploader("학생부 PDF 업로드", type=["pdf"], key="student_pdf_uploader")

if student_file and api_key:
    st.success(f"📎 학생부 파일 로드 완료: {student_file.name}")
    
    if is_unselected:
        st.warning("⚠️ 왼쪽 사이드바에서 [피드백 대분류 및 세부 평가 영역]을 먼저 선택해야 AI 평가를 시작할 수 있습니다.")
    else:
        if st.button("🔥 선택한 버전으로 AI 정밀 평가 시작하기", type="primary", use_container_width=True, key="start_eval_btn"):
            with st.spinner(f"🧠 AI 사정관이 [{evaluator_mode}] 관점에서 [{selected_feedback_type}] 맞춤 정밀 검증을 진행 중입니다..."):
                
                # RAM 상에서 읽기 (팩트 기반 텍스트 추출)
                student_text = extract_text_from_pdf_stream(student_file)
                
                if not student_text.strip():
                    st.error("❌ PDF에서 텍스트를 추출하지 못했습니다. 빈 문서이거나 이미지 스캔본인지 확인하세요.")
                    st.stop()
                    
                criteria_full_text = ""
                for fname in selected_criteria_files:
                    criteria_full_text += f"\n--- [{fname}] ---\n" + load_local_file_text(fname)[:1500]
                
                if not criteria_full_text:
                    criteria_full_text = "2028학년도 대입 표준 학종 평가 지표 적용"

                # [분기 1] 과목 세부능력 특기사항 전용 피드백
                if selected_feedback_type == "과목세부능력 특기사항 전용 피드백":
                    prompt = f"""
                    당신은 대학 입학사정관이자 교과세특 작성 컨설팅 전문가입니다.
                    제공된 [과목세특 텍스트]에는 수학(공통수학 등), 영어(공통영어 등), 통합사회, 통합과학 등 다양한 교과목의 세특 기재 내용이 포함되어 있습니다.
                    
                    [절대 준수 원칙 - 환각 방지]:
                    - 반드시 업로드된 [과목세특 텍스트]에 실제 기재된 사실과 문구만을 근거로 삼아 평가해야 합니다.
                    - 학생이 하지 않은 활동이나 생기부에 없는 내용을 AI가 스스로 지어내어 평가하거나 없는 사실을 가공해서는 절대 안 됩니다.
                    - 문서 내 특정 과목이나 내용이 다소 간결하거나 일부 누락되어 있더라도 에러를 내지 말고, 존재하는 텍스트의 맥락을 최대한 기반으로 하여 아래 7가지 채점기준(100점 만점)에 맞춰 정밀하게 평가하세요.

                    [생기부 문체 규정 - 절대 준수]:
                    - 학교생활기록부(세특)는 반드시 명사형 종결(예: ~함, ~음, ~보임, ~수행함, ~드러냄, ~탐구함)로 마무리되어야 하는 공식 문서입니다.
                    - 절대로 "-습니다", "-합니다", "-했습니다"와 같은 평서문 종결형으로 고치라는 피드백이나 수정 예시를 제시하지 마세요. 이는 생기부 작성 규정 위반이므로 명백한 오류입니다.
                    - 문장을 수정 제안할 때도 반드시 명사형 종결 어미를 유지한 채로 수정하세요.

                    [전수 점검 및 그룹핑 원칙 - 매우 중요, 절대 생략 금지]:
                    - 일부 대표 문장 1~2개만 예시로 들고 넘어가는 것을 절대 금지합니다. 업로드된 [과목세특 모음 텍스트] 전체를 처음부터 끝까지 문장 단위로 빠짐없이 검토하여, 4~6번(복붙/AI 어색함/가독성) 및 2번(교사 관찰 결여) 기준에 미달하는 "모든" 문장을 찾아내세요. 누락 없이 전수 점검하는 것이 최우선 원칙입니다.
                    - improvements와 revision_examples는 반드시 아래와 같이 "채점 항목 번호별"로 소제목을 나누어 구조화하세요:
                      [2번. 교사 관찰 결여], [4번. 복붙 의심], [5번. AI 어색함/거대 담론], [6-a. 수식어구 중첩], [6-b. 주술 호응 불일치], [6-c. 목적어 누락], [6-d. 문장 과다 길이] 등 해당되는 세부 카테고리마다 소제목을 달고, 그 아래에 해당 유형에 걸리는 모든 문제 문장을 과목명과 함께 번호를 매겨 나열하세요.
                    - 같은 항목·같은 유형 안에서 문제의 원인(왜 감점되는지)이 동일한 문장이 여러 개 발견되면, 문장마다 설명을 반복하지 말고 다음과 같이 그룹으로 묶어 요약하세요: 먼저 그 문제 유형과 이유를 한 번만 설명한 뒤, "다음 N개 문장에서 동일한 문제가 발견됨:" 형식으로 해당하는 모든 문장을 원문 그대로 나열하세요(문장마다 과목명 표기). 반대로 원인이 서로 다르면 각각 별도 항목으로 분리해서 설명하세요.
                    - revision_examples 역시 improvements에서 지적한 "모든" 문제 문장에 대해 빠짐없이 수정 전/후 예시를 제공하는 것이 원칙입니다. 다만 완전히 동일한 패턴(동일 원인)으로 그룹핑된 문장들은, 그중 1개 문장에 대해서만 상세한 '수정 전 ➔ 수정 후' 예시를 보여주고, 나머지 문장들은 "동일 유형이므로 같은 방식으로 수정 가능한 문장: 문장A, 문장B" 형태로 간략히 나열하는 방식으로 처리하세요. (문제 진단은 전수로, 수정 예시 시연은 그룹당 대표 1개로)

                    [출력 형식 규정 - 줄바꿈 필수, 절대 한 덩어리 문단 금지]:
                    - improvements와 revision_examples(그리고 good_points도 동일)는 절대로 줄바꿈 없이 이어붙인 하나의 거대한 문단으로 작성하지 마세요. JSON 문자열 값 안에 실제 줄바꿈 문자(\\n)를 반드시 포함시켜, 아래 구조를 지키세요:
                      1) 항목 소제목 "[2번. 교사 관찰 결여]" 등은 반드시 그 앞뒤로 줄바꿈을 넣어 독립된 한 줄로 작성
                      2) 그룹 설명 문장("다음 N개 문장에서 동일한 문제가 발견됨:")도 독립된 한 줄로 작성
                      3) 그 아래 나열하는 문제 문장들은 "1. (과목명/학생) '문장'" 처럼 각각 번호를 매겨 한 줄에 하나씩 줄바꿈으로 구분하여 작성 (하나의 줄에 여러 문장을 쉼표로 이어붙이지 말 것)
                      4) 서로 다른 항목([2번]과 [4번] 등) 사이에는 빈 줄을 하나 넣어 시각적으로 구분
                    - 즉, 아래 JSON 예시의 "improvements" 필드처럼 소제목-설명-번호목록이 줄 단위로 명확히 나뉜 멀티라인 문자열로 응답해야 하며, 세부 항목이 없는 good_points 역시 우수 사례를 문장별로 줄바꿈하여 나열하는 형식을 동일하게 따르세요.

                    [1~3번 항목: 근거 기반 촘촘한 피드백 원칙]:
                    - 1번(교과적 역량), 2번(교사 관찰), 3번(교과 역량) 항목은 두루뭉술한 총평이 아니라, 반드시 학생부 원문에서 실제 문장(또는 핵심 구절)을 최소 1개 이상 직접 인용하여 "왜 이 문장이 해당 항목을 충족/미충족하는지"를 구체적으로 짚어내세요.
                    - 단순히 "잘 드러남" 수준이 아니라, 어떤 개념적 깊이, 어떤 관찰 행동, 어떤 교과 고유의 사고방식이 어느 문장에서 드러나는지 문장 단위로 근거를 제시하세요.
                    - 부족한 경우, 어떤 요소(예: 교사의 1인칭 관찰 시점 부재, 개념어와 활동의 연결 고리 부족 등)가 빠졌는지 구체적으로 지적하세요.

                    [2번 항목(교사의 관찰) 판별 기준 - 매우 엄격하게 적용, 오채점 주의]:
                    - "~을 발표함", "~을 수행함", "~에 참여함", "~을 진행함", "~을 제작함"처럼 학생이 어떤 활동을 했다는 '행위 사실'만 나열된 문장은 그 자체로는 절대 좋은 '교사의 관찰'로 평가하지 마세요. 이는 활동 나열이지 관찰이 아닙니다.
                    - 진짜 '교사의 관찰'로 인정되려면, 그 활동의 과정이나 결과물에 대한 교사의 정성적 평가·판단·해석이 함께 기재되어 있어야 합니다. 예를 들어 발표 내용의 논리성/타당성/독창성에 대한 평가, 발표 태도나 질의응답 과정에서 보인 구체적 행동, 다른 학생과 차별화되는 접근 방식에 대한 언급, 교사가 관찰한 학생의 사고 과정이나 반응 등이 이에 해당합니다.
                    - 텍스트를 검토할 때 다음 두 가지를 구분하세요:
                      (1) 단순 행위 기록: "~에 대해 조사하여 발표함" → 이는 활동 사실 기록일 뿐, 2번 항목의 근거로 인정하지 마세요.
                      (2) 진짜 관찰 기록: "~에 대해 조사하여 발표함. 발표 과정에서 기존 이론의 한계를 스스로 지적하고 대안을 제시하는 등 비판적 사고력을 보임" → 밑줄 친 부분처럼 교사의 평가·해석이 붙어 있어야 인정하세요.
                    - good_points에서 2번 항목을 우수 사례로 꼽을 때는, 반드시 "교사의 평가/해석이 담긴 부분"을 정확히 지목해서 인용하세요. 세특 전체에서 '행위 나열'만 있는 문장은 모두 찾아내어(위 전수 점검 원칙에 따라) improvements에서 "활동 사실만 나열되어 있고 교사의 관찰·평가가 결여됨" 그룹으로 정리하고, 어떤 문장에 어떤 관찰·평가 내용을 보강해야 하는지 구체적으로 제안하세요.

                    [4~6번 항목: 냉정하고 철저한 정밀 검증 원칙 - 매우 중요]:
                    - 4번(복붙 여부): 세특 전체 텍스트 내에서 여러 과목 또는 여러 문단에 걸쳐 동일하거나 거의 동일한 서술어, 문장 구조, 상투적 표현(예: "~을 통해 ~을 함양함", "~에 대한 이해를 높임" 등 정형화된 틀)이 반복되는지 문장 대 문장으로 비교하여 철저히 검증하세요. 발견된 모든 유사 쌍/그룹을 빠짐없이 나열하고, 어느 문장과 어느 문장이 유사한지 구체적으로 지적하세요.
                    - 5번(AI 어색함): 구체적 활동/사례 없이 추상적이고 거대한 담론(예: "미래 사회에 기여하는 인재", "4차 산업혁명 시대") 위주로 서술되었는지, 그리고 문맥상 부자연스럽거나 과장된 수식어, 사람이 실제로 쓰지 않을 법한 어색한 한자어·기계적 번역투 표현이 있는지 문장 단위로 짚어내세요. 해당되는 모든 문장을 찾아내세요.
                    - 6번(가독성)은 아래 세부 기준으로 매우 엄격하게, 텍스트 전체에서 빠짐없이 검증하세요:
                      (a) 수식어구 중첩: 한 문장 안에 관형절이나 수식어가 2개 이상 중첩되어 핵심 내용이 바로 읽히지 않는 문장이 있는지 확인
                      (b) 주술 호응: 주어와 서술어(명사형 종결)가 문법적으로 자연스럽게 호응하는지, 주어가 바뀌었는데 서술어가 맞지 않는 경우는 없는지 확인
                      (c) 목적어 누락: 서술어가 요구하는 목적어나 대상이 생략되어 의미가 불완전한 문장이 있는지 확인
                      (d) 문장이 지나치게 길어 한 번에 읽고 이해하기 어려운 경우, 어느 지점에서 끊어야 하는지까지 구체적으로 제시
                      해당 유형 각각에 걸리는 모든 문장을 원문 그대로 인용하고, 어느 부분이 왜 문제인지 정확히 짚은 뒤, 위 [전수 점검 및 그룹핑 원칙]에 따라 정리하세요.

                    [과세특 교사 점검 채점기준 (100점 만점)]:
                    1. 학생의 교과적 역량을 잘 보여주는 기록인가 (20점 만점)
                    2. 교사의 관찰이 들어간 기록인가 (20점 만점) — 단순 활동 수행/발표 사실 나열이 아니라, 그 활동에 대한 교사의 정성적 평가·해석이 담겨 있어야 인정
                    3. 교과의 역량이 들어갔는가 (20점 만점)
                    4. 학생 간 복붙한 기록이 없는가 (10점 만점)
                    5. AI를 너무 돌려 맥락에 맞지 않는 단어나 문장이 들어가지 않았는가 (10점 만점)
                    6. 가독성이 높은가 (10점 만점)
                    7. 생기부 기재 금지 사항(대학명, 기관명, 상호명, 강사명 등)이 잘 반영되었는가 (10점 만점)

                    [업로드된 과목세특 모음 텍스트]:
                    {student_text[:8000]}

                    반드시 아래 지정된 순수 JSON 형식으로만 응답하세요 (마크다운 ```json 기호 절대 금지):
                    {{
                        "setuk_eval": {{
                            "scores": {{
                                "academic_competence": 17,
                                "teacher_observation": 16,
                                "subject_competence": 18,
                                "duplication": 9,
                                "ai_overuse": 8,
                                "readability": 8,
                                "prohibited_items": 10,
                                "total": 86
                            }},
                            "overall_summary": "수학, 영어, 통합사/과 등 과세특 기재 내용에 대한 입학사정관 관점의 팩트 기반 종합 평어",
                            "good_points": "1. (수학) '이차함수의 최댓값을 구하는 과정에서 판별식을 적용하여 풀이 과정을 논리적으로 서술함' — 개념 적용의 정확성이 돋보임\\n2. (통합과학) '실험 중 변인 통제의 필요성을 스스로 제기함' — 교사의 관찰과 평가가 함께 담긴 문장",
                            "improvements": "▍[2번. 교사 관찰 결여]\\n다음 2개 문장에서 활동 사실만 나열되고 교사의 평가·해석이 결여됨:\\n1. (수학) '이차함수 그래프를 그려 발표함.'\\n2. (영어) '팀 프로젝트에 참여함.'\\n\\n▍[6-a. 수식어구 중첩]\\n다음 2개 문장에서 수식어가 과도하게 중첩되어 핵심 내용이 바로 읽히지 않음:\\n1. (음악) '창작 과정에서 랩을 작사하는 것에 어려움을 겪기도 했으나, 모둠원들과 다각도로 아이디어를 나누고 다양한 힙합 곡들을 탐색 및 분석하여 완성도 높은 수준의 곡을 완성해 냄.'\\n2. (음악) '팀원 간 의견 조율이 필요한 상황마다 특유의 리더십과 매끄러운 진행 능력으로 분위기를 유연하게 풀어나갔으며, 가사의 분위기와 조화를 이루는 최적의 비트를 찾음.'\\n\\n(실제 응답에서는 위와 같은 형식으로 [2번]~[6-d]까지 해당되는 모든 항목을, 발견되는 모든 문제 문장을 빠짐없이 담아 작성할 것. 절대 줄바꿈 없이 이어붙이지 말 것. 단, '-습니다'체로 고치라는 피드백은 절대 포함하지 말 것)",
                            "revision_examples": "▍[2번. 교사 관찰 결여]\\n- 수정 전: '이차함수 그래프를 그려 발표함.'\\n- 수정 후: '이차함수 그래프를 그려 발표하며 그래프의 대칭성과 함수식의 관계를 스스로 도출해내는 논리적 사고력을 보임.'\\n- 동일 유형이므로 같은 방식으로 수정 가능한 문장: '팀 프로젝트에 참여함.'\\n\\n▍[6-a. 수식어구 중첩]\\n- 수정 전: '창작 과정에서 랩을 작사하는 것에 어려움을 겪기도 했으나, 모둠원들과 다각도로 아이디어를 나누고 다양한 힙합 곡들을 탐색 및 분석하여 완성도 높은 수준의 곡을 완성해 냄.'\\n- 수정 후: '창작 과정에서 랩 작사에 어려움을 겪음. 이에 모둠원들과 다각도로 아이디어를 나누고 다양한 힙합 곡들을 탐색 및 분석하며 완성도 높은 곡을 완성하는 모습을 보임.'\\n\\n(improvements에서 정리한 모든 항목·모든 그룹에 대해 위와 같은 형식으로 빠짐없이 이어서 작성할 것)"
                        }}
                    }}
                    """

                # [분기 1-2] 동아리 특기사항 전용 피드백
                elif selected_feedback_type == "동아리 특기사항 전용 피드백":
                    prompt = f"""
                    당신은 대학 입학사정관이자 동아리활동 작성 컨설팅 전문가입니다.
                    제공된 [동아리 특기사항 텍스트]는 학생들의 동아리 활동 기재 모음입니다.
                    
                    [절대 준수 원칙 - 환각 방지]:
                    - 반드시 제공된 텍스트에 있는 사실만을 근거로 평가하세요. 없는 내용을 지어내지 마세요.

                    [생기부 문체 규정 - 절대 준수]:
                    - 동아리 특기사항 역시 명사형 종결(예: ~함, ~음, ~보임, ~수행함)로 마무리되는 공식 문서입니다.
                    - 절대로 "-습니다", "-합니다"와 같은 평서문 종결형으로 고치라는 피드백이나 수정 예시를 제시하지 마세요.

                    [2번 항목(교사의 관찰) 판별 기준 - 매우 엄격하게 적용, 오채점 주의]:
                    - "~을 발표함", "~에 참여함", "~을 수행함", "~을 제작함"처럼 학생이 어떤 활동을 했다는 '행위 사실'만 나열된 문장은 그 자체로는 절대 좋은 '교사의 관찰'로 평가하지 마세요. 이는 활동 나열일 뿐 관찰이 아닙니다.
                    - 진짜 '교사의 관찰'로 인정되려면, 그 활동 과정이나 결과물에 대한 교사의 정성적 평가·해석(예: 접근 방식의 독창성, 발표·토론 과정에서 보인 태도나 사고 과정, 다른 학생과 구별되는 역할 수행 방식에 대한 평가)이 함께 기재되어 있어야 합니다.
                    - good_points에서 2번 항목을 우수 사례로 꼽을 때는 반드시 교사의 평가·해석이 담긴 부분을 정확히 인용하세요. 활동 사실만 나열된 경우라면 improvements에서 "활동 사실만 나열되고 교사의 관찰·평가가 결여됨"이라고 지적하세요.

                    [전수 점검 및 그룹핑 원칙 - 매우 중요, 절대 생략 금지]:
                    - 일부 대표 문장 1~2개만 예시로 들고 넘어가는 것을 금지합니다. [업로드된 동아리 텍스트] 전체를 문장 단위로 빠짐없이 검토하여, 2번(관찰 결여)·4번(복붙)·5번(AI 어색함)·6번(가독성: 수식어 중첩/주술 불일치/목적어 누락/과다 길이)·8번(오탈자) 기준에 미달하는 "모든" 문장을 찾아내세요.
                    - improvements와 revision_examples는 [2번. 관찰 결여], [4번. 복붙 의심], [5번. AI 어색함], [6번. 가독성], [8번. 오탈자] 등 항목별 소제목으로 구조화하고, 그 아래 해당 유형에 걸리는 모든 문제 문장을 번호를 매겨 나열하세요.
                    - 같은 항목·같은 원인으로 여러 문장이 걸릴 경우, 문장마다 반복 설명하지 말고 원인을 한 번만 설명한 뒤 "다음 N개 문장에서 동일한 문제가 발견됨:" 형식으로 해당 문장들을 나열하세요.
                    - revision_examples는 각 문제 그룹의 대표 문장 1개에 대해 상세한 '수정 전 ➔ 수정 후' 예시를 제시하고, 같은 그룹의 나머지 문장은 "동일 유형이므로 같은 방식으로 수정 가능" 형태로 간략히 안내하세요.

                    [출력 형식 규정 - 줄바꿈 필수, 절대 한 덩어리 문단 금지]:
                    - good_points, improvements, revision_examples 모두 줄바꿈 없이 이어붙인 하나의 거대한 문단으로 작성하지 마세요. JSON 문자열 값 안에 실제 줄바꿈 문자(\\n)를 반드시 포함시켜, 항목 소제목("[2번. 관찰 결여]" 등)은 독립된 한 줄로, 그룹 설명 문장도 독립된 한 줄로, 그 아래 문제 문장들은 "1. (부원명/역할) '문장'" 형태로 한 줄에 하나씩 나열하세요. 서로 다른 항목 사이에는 빈 줄을 하나 넣어 구분하세요.

                    [동아리 교사 점검 채점기준 (100점 만점)]:
                    1. 학생의 학문적 탐구 역량을 잘 보여주는 기록인가 (15점 만점)
                    2. 교사의 관찰이 들어간 기록인가 (15점 만점) — 단순 활동 수행/발표 사실 나열이 아니라, 그 활동에 대한 교사의 정성적 평가·해석이 담겨 있어야 인정
                    3. 학생이 이수한 교과목에서 심화 탐구한 내용이 들어갔는가 (15점 만점)
                    4. 학생 간 복붙한 기록이 없는가 (10점 만점)
                    5. AI를 너무 돌려 맥락에 맞지 않는 단어나 문장이 들어가지 않았는가 (10점 만점)
                    6. 가독성이 높은가 (15점 만점)
                    7. 생기부 기재 금지 사항이 잘 반영되어 있는가 (10점 만점)
                    8. 오탈자가 없는가 (10점 만점)

                    [업로드된 동아리 텍스트]:
                    {student_text[:8000]}

                    반드시 아래 지정된 순수 JSON 형식으로만 응답하세요 (마크다운 ```json 기호 절대 금지):
                    {{
                        "club_eval": {{
                            "scores": {{
                                "academic_inquiry": 13,
                                "teacher_observation": 13,
                                "subject_connection": 13,
                                "duplication": 9,
                                "ai_overuse": 9,
                                "readability": 13,
                                "prohibited_items": 10,
                                "typo_check": 10,
                                "total": 90
                            }},
                            "overall_summary": "동아리 특기사항에 대한 입학사정관 관점의 팩트 기반 종합 평어",
                            "good_points": "1. (김OO) '토론 진행 중 상반된 의견을 정리하여 절충안을 제시함' — 리더십과 조율 능력에 대한 교사의 평가가 담긴 문장\\n2. (이OO) '데이터 분석 결과의 한계를 스스로 지적함' — 비판적 사고에 대한 관찰이 드러남",
                            "improvements": "▍[2번. 관찰 결여]\\n다음 2개 문장에서 활동 사실만 나열되고 교사의 평가·해석이 결여됨:\\n1. (박OO) '토론에 참여함.'\\n2. (최OO) '보고서를 작성함.'\\n\\n▍[6번. 가독성]\\n다음 문장에서 수식어가 과도하게 중첩되어 가독성이 저하됨:\\n1. (정OO) '팀원 간 의견 조율이 필요한 상황마다 특유의 리더십과 매끄러운 진행 능력으로 분위기를 유연하게 풀어나갔으며, 가사의 분위기와 조화를 이루는 최적의 비트를 찾음.'\\n\\n(실제 응답에서는 위와 같은 형식으로 해당되는 모든 항목에 대해, 발견되는 모든 문제 문장을 빠짐없이 담아 작성할 것. 단, '-습니다'체로 고치라는 피드백은 절대 포함하지 말 것)",
                            "revision_examples": "▍[2번. 관찰 결여]\\n- 수정 전: '토론에 참여함.'\\n- 수정 후: '토론에 참여하여 상반된 의견 사이에서 절충안을 제시하는 조율 능력을 보임.'\\n- 동일 유형이므로 같은 방식으로 수정 가능한 문장: '보고서를 작성함.'\\n\\n(improvements에서 정리한 모든 항목·모든 그룹에 대해 위와 같은 형식으로 빠짐없이 이어서 작성할 것)"
                        }}
                    }}
                    """

                # [분기 2] 교사 전용 (생기부 종합) 피드백 (교사 기록 적절성, 가독성, 성적 수준 적절성, AI 남용 여부 및 향후 지도 방향 포함)
                elif feedback_category == "교사전용 피드백 버전":
                    prompt = f"""
                    당신은 전국 대학부종합전형 서류를 평가하는 [{evaluator_mode}]입니다.
                    제공된 [학생 생기부 텍스트]에는 수학, 영어, 통합사회, 통합과학 등 다양한 교과 기록이 포함되어 있습니다.
                    
                    [절대 준수 원칙 - 환각 방지]:
                    - 반드시 생기부 PDF에 실제로 적힌 텍스트 내용만을 근거로 평가하세요. 없는 내용을 만들어내지 마세요.

                    [핵심 평가 지침 (생기부 종합 전용 피드백)]:
                    1. 평가자 관점: '{evaluator_mode}' 특성 반영.
                    2. 학업(40점)/진로(40점)/공동체(20점) 총 100점 만점으로 점수를 매기세요.
                    3. 학생 성적과 활동에 대한 교사의 기록이 적절했는지 상세히 진단하세요:
                       - 교사의 관찰(구체적 행동이나 수업 태도 등)이 잘 드러났는지. 단, "~을 발표함", "~에 참여함", "~을 수행함"처럼 활동을 했다는 사실만 나열된 문장은 좋은 관찰 기록으로 인정하지 마세요. 그 활동의 과정이나 결과물에 대한 교사의 정성적 평가·해석(태도, 사고 과정, 접근 방식의 특징 등)이 함께 담겨 있어야 진짜 관찰 기록으로 판단하세요.
                       - 기록의 가독성과 문장 구조가 매끄러운지
                       - 학생의 성적 수준(원점수, 과목 평균, 등급 등)에 어울리는 기록인지
                       - AI를 남용한 티(거대 담론 나열, 구체적 사례 부재 등)가 나는지 여부
                    4. 학생에게 앞으로 어떤 방향으로 활동을 보완하고 확장해야 할지 구체적인 지도 방향과 피드백을 제시하세요.

                    [학생 제출 텍스트]:
                    {student_text[:8000]}

                    반드시 아래 지정된 순수 JSON 형식으로만 응답하세요 (마크다운 ```json 기호 절대 금지):
                    {{
                        "scores": {{
                            "academic": 33.0,
                            "career": 32.5,
                            "community": 16.5,
                            "total": 82.0
                        }},
                        "teacher_feedback": {{
                            "category": "{selected_feedback_type}",
                            "strength": "교사 기록의 적절성(관찰, 가독성, 성적 수준 부합도 등) 측면에서의 우수한 장점 서술",
                            "weakness": "교사 기록의 보완점(AI 남용 흔적, 구체적 관찰 부족 등) 및 감점 사유 지적, 그리고 앞으로 학생을 어떤 방향으로 지도해야 할지에 대한 구체적 제언",
                            "quote": "텍스트에서 실제 인용한 핵심 문장"
                        }}
                    }}
                    """

                # [분기 3] 학생 전용 피드백
                else:
                    prompt = f"""
                    당신은 전국 대학부종합전형 서류를 평가하는 [{evaluator_mode}]입니다.
                    제공된 [학생 생기부 텍스트]에는 공통수학, 공통영어, 통합사회, 통합과학 등 다양한 과목 기록이 포함되어 있습니다.
                    
                    [절대 준수 원칙 - 환각 방지]:
                    - 반드시 실제 기재된 생기부 텍스트 내용만을 근거로 진단하세요. 지어낸 활동으로 평가하지 마세요.

                    [핵심 평가 지침]:
                    1. 평가자 관점: '{evaluator_mode}' 특성 반영.
                    2. 학업(40점)/진로(40점)/공동체(20점) 총 100점 만점으로 점수를 산출하세요.
                    3. 학생용 피드백: 헛된 희망을 주지 않는 입학사정관 관점의 냉정한 현재 위치 진단(지원 가능 대학 라인), 활동 강점과 치명적 약점, 앞으로 실행해야 할 구체적인 탐구 주제 및 활동 솔루션을 제시하세요.

                    [학생 제출 텍스트]:
                    {student_text[:8000]}

                    반드시 아래 지정된 순수 JSON 형식으로만 응답하세요 (마크다운 ```json 기호 절대 금지):
                    {{
                        "scores": {{
                            "academic": 33.0,
                            "career": 32.5,
                            "community": 16.5,
                            "total": 82.0
                        }},
                        "student_feedback": {{
                            "current_position": "입학사정관 관점 냉정한 현위치 진단 및 지원 가능 대학 라인",
                            "strength_analysis": "여태까지 한 활동의 핵심 강점 분석",
                            "weakness_analysis": "치명적인 약점 및 감점 요소 분석",
                            "recommendation": "앞으로 실행해야 할 구체적인 탐구 주제 및 과목 선택/활동 솔루션"
                        }}
                    }}
                    """

                try:
                    model = genai.GenerativeModel(model_option)
                    response = model.generate_content(prompt)
                    cleaned = response.text.strip().replace("```json", "").replace("```", "").strip()
                    result_json = json.loads(cleaned)

                    # AI가 줄바꿈 없는 한 덩어리 문단으로 응답하는 경우를 대비한 코드 단 정규화
                    _TEXT_FIELDS = ["overall_summary", "good_points", "improvements", "revision_examples",
                                     "strength", "weakness", "quote",
                                     "current_position", "strength_analysis", "weakness_analysis", "recommendation"]
                    for _outer_key in ("setuk_eval", "club_eval"):
                        if _outer_key in result_json and isinstance(result_json[_outer_key], dict):
                            for _field in _TEXT_FIELDS:
                                if _field in result_json[_outer_key]:
                                    result_json[_outer_key][_field] = normalize_feedback_text(result_json[_outer_key][_field])
                    for _sub_key in ("teacher_feedback", "student_feedback"):
                        if _sub_key in result_json and isinstance(result_json[_sub_key], dict):
                            for _field in _TEXT_FIELDS:
                                if _field in result_json[_sub_key]:
                                    result_json[_sub_key][_field] = normalize_feedback_text(result_json[_sub_key][_field])

                    st.session_state["eval_result"] = result_json
                    st.session_state["evaluated_filename"] = student_file.name
                    st.session_state["eval_mode_title"] = selected_feedback_type
                    st.success("🎉 분석 완료! 제출된 학생부 데이터는 메모리(RAM)에서 즉시 영구 파기되었습니다.")
                except json.JSONDecodeError:
                    st.error("⚠️ AI 응답 형식을 해석하는 데 실패했습니다. 한 번 더 실행하시거나 모델을 변경해 보세요.")
                    st.code(response.text)
                except Exception as e:
                    st.error(f"평가 중 오류가 발생했습니다: {e}")

# --- 8. 채점 결과 및 맞춤형 피드백 출력 구역 ---
if "eval_result" in st.session_state:
    res = st.session_state["eval_result"]
    fb_title = st.session_state.get("eval_mode_title", "")
    
    st.divider()
    
    # 📌 과목 세특 교사 점검 모드 결과 출력
    if fb_title == "과목세부능력 특기사항 전용 피드백" and "setuk_eval" in res:
        st_eval = res["setuk_eval"]
        scores = st_eval.get("scores", {})
        
        st.markdown(f"### 📊 과목세부능력 특기사항 교사 기재 점검 표 (`{evaluator_mode}` 기준)")
        st.metric("✨ 과세특 기재 품질 총점", f"{scores.get('total', 0)} / 100 점")
        
        st.divider()
        st.markdown("#### 📋 7대 세부 점검 항목별 점수")
        sc1, sc2, sc3, sc4 = st.columns(4)
        sc1.metric("1. 학생 교과역량", f"{scores.get('academic_competence', 0)} / 20점")
        sc2.metric("2. 교사 직접 관찰", f"{scores.get('teacher_observation', 0)} / 20점")
        sc3.metric("3. 교과 핵심 역량", f"{scores.get('subject_competence', 0)} / 20점")
        sc4.metric("4. 학생간 복붙 방지", f"{scores.get('duplication', 0)} / 10점")
        
        sc5, sc6, sc7, _ = st.columns(4)
        sc5.metric("5. AI 대필 오염 방지", f"{scores.get('ai_overuse', 0)} / 10점")
        sc6.metric("6. 가독성 및 문장", f"{scores.get('readability', 0)} / 10점")
        sc7.metric("7. 기재금지 사항 준수", f"{scores.get('prohibited_items', 0)} / 10점")
        
        st.divider()
        st.subheader("👩‍🏫 과세특 교사 자가점검 및 개선 피드백")
        st.info(f"**📝 종합 총평:** {st_eval.get('overall_summary', '')}")
        st.success(f"**👍 기재 우수 사항:** {st_eval.get('good_points', '')}")
        st.warning(f"**⚠️ 보완 및 수정 필요사항:** {st_eval.get('improvements', '')}")
        st.error(f"**✏️ 수정·보완 추천 문장 예시:**\n\n{st_eval.get('revision_examples', '')}")

    # 📌 동아리 교사 점검 모드 결과 출력
    elif fb_title == "동아리 특기사항 전용 피드백" and "club_eval" in res:
        c_eval = res["club_eval"]
        c_scores = c_eval.get("scores", {})
        
        st.markdown(f"### 📊 동아리 특기사항 교사 기재 점검 표 (`{evaluator_mode}` 기준)")
        st.metric("✨ 동아리 기재 품질 총점", f"{c_scores.get('total', 0)} / 100 점")
        
        st.divider()
        st.markdown("#### 📋 8대 세부 점검 항목별 점수")
        cc1, cc2, cc3, cc4 = st.columns(4)
        cc1.metric("1. 학문적 탐구 역량", f"{c_scores.get('academic_inquiry', 0)} / 15점")
        cc2.metric("2. 교사 직접 관찰", f"{c_scores.get('teacher_observation', 0)} / 15점")
        cc3.metric("3. 교과 심화 연계", f"{c_scores.get('subject_connection', 0)} / 15점")
        cc4.metric("4. 복붙 방지", f"{c_scores.get('duplication', 0)} / 10점")
        
        cc5, cc6, cc7, cc8 = st.columns(4)
        cc5.metric("5. AI 오염 방지", f"{c_scores.get('ai_overuse', 0)} / 10점")
        cc6.metric("6. 가독성", f"{c_scores.get('readability', 0)} / 15점")
        cc7.metric("7. 기재금지 준수", f"{c_scores.get('prohibited_items', 0)} / 10점")
        cc8.metric("8. 오탈자 검수", f"{c_scores.get('typo_check', 0)} / 10점")
        
        st.divider()
        st.subheader("👩‍🏫 동아리 교사 자가점검 및 개선 피드백")
        st.info(f"**📝 종합 총평:** {c_eval.get('overall_summary', '')}")
        st.success(f"**👍 기재 우수 사항:** {c_eval.get('good_points', '')}")
        st.warning(f"**⚠️ 보완 및 수정 필요사항:** {c_eval.get('improvements', '')}")
        st.error(f"**✏️ 수정·보완 추천 문장 예시:**\n\n{c_eval.get('revision_examples', '')}")

    # 📌 표준 생기부 / 기타 교사 / 학생 피드백 출력
    else:
        scores = res.get("scores", {})
        st.markdown(f"### 📊 평가 스코어 카드 (`{evaluator_mode}` 관점 / `{fb_title}`)")
        
        col_s1, col_s2, col_s3, col_s4 = st.columns(4)
        col_s1.metric("📕 학업역량", f"{scores.get('academic', 0)} / 40 점")
        col_s2.metric("📗 진로역량", f"{scores.get('career', 0)} / 40 점")
        col_s3.metric("📘 공동체역량", f"{scores.get('community', 0)} / 20 점")
        col_s4.metric("✨ 최종 종합 점수", f"{scores.get('total', 0)} / 100 점")
        
        st.divider()
        
        if "teacher_feedback" in res:
            st.subheader(f"👩‍🏫 NEIS 및 진학 지도용 교사전용 피드백 (`{fb_title}`)")
            tf = res["teacher_feedback"]
            st.success(f"**👍 핵심 장점:** {tf.get('strength', '')}")
            st.warning(f"**⚠️ 보완점 및 감점 사유:** {tf.get('weakness', '')}")
            if tf.get('quote'):
                st.info(f"**🎯 원문 인용 근거:** \"{tf.get('quote')}\"")

        if "student_feedback" in res:
            st.subheader("🎓 학생 전용 쓴소리 진단 및 탐구 솔루션 리포트")
            sf = res["student_feedback"]
            
            st.markdown("#### 🔍 1. 입학사정관 관점 냉정한 현재 위치 (지원 가능 대학 라인)")
            st.error(sf.get("current_position", ""))
            
            col_st1, col_st2 = st.columns(2)
            with col_st1:
                st.markdown("#### 👍 2. 기존 활동의 주요 강점")
                st.success(sf.get("strength_analysis", ""))
            with col_st2:
                st.markdown("#### 🚨 3. 치명적인 약점 및 감점 요소")
                st.warning(sf.get("weakness_analysis", ""))
                
            st.markdown("#### 🚀 4. 앞으로의 구체적 추천 활동 및 탐구 주제 솔루션")
            st.info(sf.get("recommendation", ""))

    st.divider()
    
    # --- 9. PDF / DOCX 다운로드 기능 (요구사항 5) ---
    st.markdown("### 📥 3단계: 정밀 진단 보고서 다운로드")
    d1, d2 = st.columns(2)
    
    with d1:
        if REPORTLAB_AVAILABLE:
            pdf_bytes = generate_pdf_report(
                res, 
                st.session_state.get("evaluated_filename", "student.pdf"), 
                evaluator_mode, 
                fb_title
            )
            st.download_button(
                label="📄 정밀 진단 보고서 PDF 다운로드",
                data=pdf_bytes,
                file_name=f"브니엘고_AI_생기부평가_{evaluator_mode}_{st.session_state.get('evaluated_filename', 'student').replace('.pdf','')}.pdf",
                mime="application/pdf",
                use_container_width=True,
                key="download_pdf_btn"
            )
        else:
            st.warning("ReportLab 모듈이 누락되었습니다. (`pip install reportlab` 필요)")
            
    with d2:
        if DOCX_AVAILABLE:
            docx_bytes = generate_docx_report(
                res, 
                st.session_state.get("evaluated_filename", "student.pdf"), 
                evaluator_mode, 
                fb_title
            )
            st.download_button(
                label="📝 정밀 진단 보고서 WORD(DOCX) 다운로드",
                data=docx_bytes,
                file_name=f"브니엘고_AI_생기부평가_{evaluator_mode}_{st.session_state.get('evaluated_filename', 'student').replace('.pdf','')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key="download_docx_btn"
            )
        else:
            st.warning("python-docx 모듈이 누락되었습니다. (`pip install python-docx` 필요)")
